"""
Input normalization utilities.

All processing engines operate on .docx files. This module accepts supported
input formats and converts them into temporary .docx files when needed.
"""

from __future__ import annotations

import io
import os
from pathlib import Path

from docx import Document
from docx.shared import Inches

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - handled at runtime
    PdfReader = None

try:
    from PIL import Image
except Exception:  # pragma: no cover - handled at runtime
    Image = None

try:
    import pytesseract
except Exception:  # pragma: no cover - handled at runtime
    pytesseract = None

try:
    import fitz
except Exception:  # pragma: no cover - handled at runtime
    fitz = None


SUPPORTED_EXTENSIONS = {
    ".docx",
    ".pdf",
    ".png",
    ".jpg",
    ".jpeg",
    ".bmp",
    ".gif",
    ".tif",
    ".tiff",
    ".webp",
}


def is_supported_filename(filename: str) -> bool:
    return Path(filename or "").suffix.lower() in SUPPORTED_EXTENSIONS


def supported_extensions_csv() -> str:
    """Return a stable comma-separated accept string for HTML/file validators."""
    order = (
        ".docx",
        ".pdf",
        ".png",
        ".jpg",
        ".jpeg",
        ".bmp",
        ".gif",
        ".tif",
        ".tiff",
        ".webp",
    )
    return ",".join(order)


def normalize_input_to_docx(input_path: str, output_docx_path: str) -> dict:
    """
    Normalize supported file types into a .docx input.

    Returns:
        {
            "path": str,      # path to the .docx used for processing
            "converted": bool,
            "note": str | None,
        }
    """
    ext = Path(input_path).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: {supported_extensions_csv()}"
        )

    if ext == ".docx":
        return {"path": input_path, "converted": False, "note": None}

    if ext == ".pdf":
        note = _pdf_to_docx(input_path, output_docx_path)
        return {"path": output_docx_path, "converted": True, "note": note}

    note = _image_to_docx(input_path, output_docx_path)
    return {"path": output_docx_path, "converted": True, "note": note}


def _pdf_to_docx(input_path: str, output_docx_path: str) -> str:
    doc = Document()
    doc.add_heading("Imported PDF Content", level=1)
    doc.add_paragraph(f"Source file: {os.path.basename(input_path)}")
    doc.add_paragraph("")

    if PdfReader is None and not _can_ocr_pdf():
        doc.add_paragraph(
            "PDF text extraction is unavailable because required dependencies are missing."
        )
        doc.save(output_docx_path)
        return "PDF uploaded, but text extraction was unavailable (missing pypdf/OCR stack)."

    reader = PdfReader(input_path) if PdfReader is not None else None
    pdf_doc = fitz.open(input_path) if _can_ocr_pdf() else None
    extracted_pages = 0
    ocr_pages = 0

    try:
        page_count = _resolve_page_count(reader, pdf_doc)
        for page_idx in range(page_count):
            page_num = page_idx + 1
            text = ""
            if reader is not None:
                try:
                    text = (reader.pages[page_idx].extract_text() or "").strip()
                except Exception:
                    text = ""

            used_ocr = False
            if not text and pdf_doc is not None:
                text = _ocr_pdf_page(pdf_doc[page_idx])
                used_ocr = bool(text)

            text = (text or "").strip()
            if not text:
                continue

            extracted_pages += 1
            if used_ocr:
                ocr_pages += 1
            if extracted_pages > 1:
                doc.add_page_break()
            doc.add_heading(f"Page {page_num}", level=2)
            _append_text(doc, text)
    finally:
        if pdf_doc is not None:
            pdf_doc.close()

    if extracted_pages == 0:
        doc.add_paragraph(
            "No extractable text was found in the PDF, including OCR attempts."
        )
        note = (
            "PDF converted with no extractable text (OCR did not return usable content)."
        )
    elif ocr_pages > 0:
        note = (
            f"PDF converted to DOCX with text on {extracted_pages} page(s), "
            f"including OCR on {ocr_pages} scanned page(s)."
        )
    else:
        note = f"PDF converted to DOCX using extracted text from {extracted_pages} page(s)."

    doc.save(output_docx_path)
    return note


def _image_to_docx(input_path: str, output_docx_path: str) -> str:
    doc = Document()
    doc.add_heading("Imported Image Content", level=1)
    doc.add_paragraph(f"Source file: {os.path.basename(input_path)}")
    doc.add_paragraph("")

    try:
        doc.add_picture(input_path, width=Inches(6.0))
    except Exception as exc:
        doc.add_paragraph(f"Unable to embed image preview: {exc}")

    ocr_text = _ocr_image_path(input_path)
    if ocr_text:
        doc.add_heading("OCR Extracted Text", level=2)
        _append_text(doc, ocr_text)
        note = "Image converted to DOCX with OCR text extraction."
    elif _can_ocr_image():
        doc.add_paragraph("OCR ran but did not return readable text.")
        note = "Image converted to DOCX (OCR attempted but no text detected)."
    else:
        doc.add_paragraph(
            "OCR is unavailable. Install pytesseract and the Tesseract binary "
            "to extract text from image uploads."
        )
        note = "Image converted to DOCX (embedded preview, OCR unavailable)."

    doc.save(output_docx_path)
    return note


def _resolve_page_count(reader, pdf_doc) -> int:
    if reader is not None:
        return len(reader.pages)
    if pdf_doc is not None:
        return pdf_doc.page_count
    return 0


def _append_text(doc: Document, text: str) -> None:
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        doc.add_paragraph(line)


def _can_ocr_image() -> bool:
    return bool(Image is not None and pytesseract is not None)


def _can_ocr_pdf() -> bool:
    return bool(fitz is not None and _can_ocr_image())


def _ocr_pdf_page(page) -> str:
    if not _can_ocr_pdf():
        return ""
    try:
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        with Image.open(io.BytesIO(pix.tobytes("png"))) as img:
            return (pytesseract.image_to_string(img) or "").strip()
    except Exception:
        return ""


def _ocr_image_path(image_path: str) -> str:
    if not _can_ocr_image():
        return ""
    try:
        with Image.open(image_path) as img:
            return (pytesseract.image_to_string(img) or "").strip()
    except Exception:
        return ""
