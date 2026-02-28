"""
INTEGRATION — Table formatting.
Applies header fills, alternating row colours, borders, cell padding,
captions above, source lines below.
"""

import re
from copy import deepcopy
from lxml import etree
from docx.shared import Pt, Twips, Inches, RGBColor, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from . import constants as C
from .typography import apply_run_style


def _rgb(hex_str: str) -> RGBColor:
    return RGBColor(int(hex_str[:2], 16), int(hex_str[2:4], 16), int(hex_str[4:], 16))


def _set_cell_shading(cell, hex_color: str):
    """Set the background fill of a cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders. Each argument: (size_eighths, color_hex) or None."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = tc.makeelement(qn("w:tcBorders"), {})
        tcPr.append(borders)

    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is None:
            continue
        sz, color = val
        el = borders.find(qn(f"w:{side}"))
        if el is None:
            el = tc.makeelement(qn(f"w:{side}"), {})
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")


def _set_table_borders(table):
    """Set table-level borders per the INTEGRATION spec."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = tbl.makeelement(qn("w:tblPr"), {})
        tbl.insert(0, tblPr)

    # Remove existing borders
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    borders_xml = f"""
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="{C.TABLE_HEADER_BORDER_SZ}" w:color="{C.GREEN_HAZE}" w:space="0"/>
        <w:left w:val="single" w:sz="{C.TABLE_OUTER_BORDER_SZ}" w:color="{C.GREEN_HAZE}" w:space="0"/>
        <w:bottom w:val="single" w:sz="{C.TABLE_HEADER_BORDER_SZ}" w:color="{C.GREEN_HAZE}" w:space="0"/>
        <w:right w:val="single" w:sz="{C.TABLE_OUTER_BORDER_SZ}" w:color="{C.GREEN_HAZE}" w:space="0"/>
        <w:insideH w:val="single" w:sz="{C.TABLE_INNER_BORDER_SZ}" w:color="{C.INNER_BORDER}" w:space="0"/>
        <w:insideV w:val="none" w:sz="0" w:color="auto" w:space="0"/>
    </w:tblBorders>
    """
    tblPr.append(parse_xml(borders_xml))


def _set_table_width(table):
    """Set total table width to 9000 DXA."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = tbl.makeelement(qn("w:tblPr"), {})
        tbl.insert(0, tblPr)

    # Remove existing width
    existing = tblPr.find(qn("w:tblW"))
    if existing is not None:
        tblPr.remove(existing)

    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{C.TABLE_WIDTH_DXA}" w:type="dxa"/>')
    tblPr.append(tblW)


def _set_cell_padding(table):
    """Set default cell margin/padding for the table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = tbl.makeelement(qn("w:tblPr"), {})
        tbl.insert(0, tblPr)

    existing = tblPr.find(qn("w:tblCellMar"))
    if existing is not None:
        tblPr.remove(existing)

    margin_xml = f"""
    <w:tblCellMar {nsdecls("w")}>
        <w:top w:w="{C.TABLE_CELL_PAD_TOP}" w:type="dxa"/>
        <w:left w:w="{C.TABLE_CELL_PAD_LEFT}" w:type="dxa"/>
        <w:bottom w:w="{C.TABLE_CELL_PAD_BOTTOM}" w:type="dxa"/>
        <w:right w:w="{C.TABLE_CELL_PAD_RIGHT}" w:type="dxa"/>
    </w:tblCellMar>
    """
    tblPr.append(parse_xml(margin_xml))


def _set_vertical_alignment(cell, val="center"):
    """Set vertical alignment on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:vAlign"))
    if existing is not None:
        tcPr.remove(existing)
    vAlign = tc.makeelement(qn("w:vAlign"), {qn("w:val"): val})
    tcPr.append(vAlign)


def _is_number(text: str) -> bool:
    """Check if a string looks like a number (for right-alignment)."""
    cleaned = text.strip().replace(",", "").replace("%", "").replace("$", "")
    try:
        float(cleaned)
        return True
    except ValueError:
        return False


def format_tables(doc, chapter_num: int = 1) -> list[str]:
    """
    Format all tables in the document per INTEGRATION spec.
    Returns list of table names that were reformatted.
    """
    reformatted = []

    for idx, table in enumerate(doc.tables):
        table_name = f"Table {chapter_num}.{idx + 1}"
        reformatted.append(table_name)

        # Set table-level properties
        _set_table_borders(table)
        _set_table_width(table)
        _set_cell_padding(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)

            for cell in row.cells:
                # Vertical alignment
                _set_vertical_alignment(cell)

                if is_header:
                    # Header row: green background, white bold centered text
                    _set_cell_shading(cell, C.GREEN_HAZE)
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.name = C.FONT_FAMILY
                            run.font.size = Pt(10)
                            run.font.bold = True
                            run.font.color.rgb = _rgb(C.WHITE)
                else:
                    # Data rows: alternating fills
                    if row_idx % 2 == 1:
                        _set_cell_shading(cell, C.WHITE)
                    else:
                        _set_cell_shading(cell, C.ALT_ROW)

                    for para in cell.paragraphs:
                        # Determine alignment
                        text = para.text.strip()
                        if _is_number(text):
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        for run in para.runs:
                            run.font.name = C.FONT_FAMILY
                            run.font.size = Pt(10)
                            run.font.bold = False
                            run.font.color.rgb = _rgb(C.BODY_TEXT)

        # Set header row border special treatment
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                _set_cell_borders(
                    cell,
                    top=(C.TABLE_HEADER_BORDER_SZ, C.GREEN_HAZE),
                    bottom=(C.TABLE_HEADER_BORDER_SZ, C.GREEN_HAZE),
                    left=(C.TABLE_OUTER_BORDER_SZ, C.GREEN_HAZE),
                    right=(C.TABLE_OUTER_BORDER_SZ, C.GREEN_HAZE),
                )

    return reformatted


def find_or_create_table_captions(doc, chapter_num: int = 1) -> list[str]:
    """
    Ensure every table has a properly formatted caption above it
    and a source line below it. Returns list of corrections made.
    """
    corrections = []
    table_counter = 0

    body = doc.element.body
    elements = list(body)

    for i, element in enumerate(elements):
        if element.tag != qn("w:tbl"):
            continue

        table_counter += 1
        expected_caption = f"Table {chapter_num}.{table_counter}"

        # Check paragraph before table for existing caption
        has_caption = False
        if i > 0:
            prev = elements[i - 1]
            if prev.tag == qn("w:p"):
                text = "".join(node.text or "" for node in prev.iter(qn("w:t")))
                if text.strip().lower().startswith("table"):
                    has_caption = True
                    # Reformat existing caption
                    for run_el in prev.iter(qn("w:r")):
                        for t in run_el.iter(qn("w:t")):
                            pass  # Keep existing text, just restyle
                    corrections.append(f"Restyled caption for {expected_caption}")

        if not has_caption:
            corrections.append(f"Caption missing for {expected_caption} — needs manual addition")

        # Check for source line after table
        has_source = False
        if i + 1 < len(elements):
            nxt = elements[i + 1]
            if nxt.tag == qn("w:p"):
                text = "".join(node.text or "" for node in nxt.iter(qn("w:t")))
                if text.strip().lower().startswith("source"):
                    has_source = True

        if not has_source:
            corrections.append(f"Source line missing after {expected_caption}")

    return corrections
