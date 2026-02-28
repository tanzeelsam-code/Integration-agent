"""
INTEGRATION — Core formatting engine (orchestrator).
Loads a .docx, runs all 12 processing steps in order,
returns the document and a compliance summary.
"""

from docx import Document
from docx.shared import Pt, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from . import constants as C
from .typography import enforce_arial_everywhere, apply_paragraph_style
from .headings import renumber_headings
from .tables import format_tables, find_or_create_table_captions
from .figures import format_figures
from .layout import apply_page_layout, apply_header, apply_footer
from .lists import format_bullet_lists
from .compliance import generate_summary


def _rgb(hex_str: str) -> RGBColor:
    return RGBColor(int(hex_str[:2], 16), int(hex_str[2:4], 16), int(hex_str[4:], 16))


def _detect_chapter(doc) -> int:
    """Try to detect the current chapter number from the first H1."""
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name == "Heading 1":
            text = para.text.strip()
            if text and text[0].isdigit():
                try:
                    return int(text.split(".")[0])
                except ValueError:
                    pass
            return 1
    return 1


def _fix_body_spacing(doc) -> int:
    """Apply 6pt before / 6pt after + single (1.08) line spacing to body paragraphs."""
    count = 0
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        # Skip headings, captions, source lines, etc.
        if style_name.startswith("Heading"):
            continue
        if style_name in ("Caption", "Footnote Text"):
            continue

        pf = para.paragraph_format
        pf.space_before = Twips(C.BODY_SPACE_BEFORE)
        pf.space_after = Twips(C.BODY_SPACE_AFTER)

        # Line spacing: 1.08 (single)
        pf.line_spacing = C.LINE_SPACING

        # Justify body paragraphs
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        count += 1

    return count


def _fix_hyperlinks(doc) -> int:
    """Style all hyperlinks as blue #0563C1, underlined."""
    count = 0
    target_rgb = _rgb(C.HYPERLINK)

    for para in doc.paragraphs:
        # Check for hyperlink elements in the XML
        hyperlinks = para._element.findall(qn("w:hyperlink"))
        for hl in hyperlinks:
            for run_el in hl.findall(qn("w:r")):
                rPr = run_el.find(qn("w:rPr"))
                if rPr is None:
                    rPr = run_el.makeelement(qn("w:rPr"), {})
                    run_el.insert(0, rPr)

                # Set colour
                color_el = rPr.find(qn("w:color"))
                if color_el is None:
                    color_el = run_el.makeelement(qn("w:color"), {})
                    rPr.append(color_el)
                color_el.set(qn("w:val"), C.HYPERLINK)

                # Set underline
                u_el = rPr.find(qn("w:u"))
                if u_el is None:
                    u_el = run_el.makeelement(qn("w:u"), {})
                    rPr.append(u_el)
                u_el.set(qn("w:val"), "single")

                count += 1

    return count


def _remove_noncompliant_colours(doc) -> list[str]:
    """
    Scan all runs for non-compliant colours and replace them with
    the appropriate compliant colour (usually BODY_TEXT black).
    """
    allowed = {
        C.GREEN_HAZE.upper(), C.WHITE.upper(), C.ALT_ROW.upper(),
        C.BODY_TEXT.upper(), C.SOURCE_TEXT.upper(), C.HEADER_GREY.upper(),
        C.INNER_BORDER.upper(), C.HYPERLINK.upper(),
    }
    corrections = []

    for para in doc.paragraphs:
        # Skip headings — they use GREEN_HAZE which is allowed
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            continue

        for run in para.runs:
            if run.font.color and run.font.color.rgb:
                hex_val = str(run.font.color.rgb).upper()
                if hex_val not in allowed:
                    corrections.append(f"#{hex_val}")
                    # Default to body text colour
                    run.font.color.rgb = _rgb(C.BODY_TEXT)

    return corrections


def reformat_document(
    input_path: str,
    output_path: str,
    report_name: str = "Report",
    year: str = "2026",
    project_number: str = "PRJ-001",
) -> str:
    """
    Main entry point. Reformats a .docx file per INTEGRATION spec.

    Args:
        input_path: Path to input .docx file
        output_path: Path to save reformatted .docx
        report_name: Name for the header template
        year: Year for footer copyright
        project_number: Project number for footer

    Returns:
        Compliance summary string
    """
    doc = Document(input_path)
    chapter_num = _detect_chapter(doc)

    # ── Step 1: Scan & classify (implicit — done by each module) ──

    # ── Step 2: Apply typography (Arial everywhere) — BEFORE headings
    #    so heading styles applied later take precedence
    fonts_corrected = enforce_arial_everywhere(doc)

    # ── Step 3: Renumber headings (applies bold + green color) ──
    headings_changed = renumber_headings(doc)

    # ── Step 4: Fix spacing + justify body paragraphs ──
    spacing_count = _fix_body_spacing(doc)

    # ── Step 5: Reformat tables ──
    tables_reformatted = format_tables(doc, chapter_num=chapter_num)

    # ── Step 6: Reformat figures ──
    figures_info = format_figures(doc, chapter_num=chapter_num)

    # ── Step 7: Renumber tables & figures (done within steps 5 & 6) ──
    caption_corrections = find_or_create_table_captions(doc, chapter_num=chapter_num)

    # ── Step 8: Apply header/footer ──
    header_ok = apply_header(doc, report_name=report_name)
    footer_ok = apply_footer(doc, year=year, project_number=project_number)

    # ── Step 9: Apply page layout ──
    apply_page_layout(doc)

    # ── Step 10: Bullet lists ──
    bullets_fixed = format_bullet_lists(doc)

    # ── Step 11: Hyperlinks ──
    hyperlinks_fixed = _fix_hyperlinks(doc)

    # ── Step 12: Remove non-compliant colours ──
    colours_corrected = _remove_noncompliant_colours(doc)

    # ── Save ──
    doc.save(output_path)

    # ── Generate compliance summary ──
    summary = generate_summary(
        headings_changed=headings_changed,
        tables_reformatted=tables_reformatted,
        figures_info=figures_info,
        caption_corrections=caption_corrections,
        fonts_corrected=fonts_corrected,
        colours_corrected=colours_corrected,
        spacing_corrected=spacing_count,
        header_applied=header_ok,
        footer_applied=footer_ok,
        bullets_fixed=bullets_fixed,
        hyperlinks_fixed=hyperlinks_fixed,
    )

    return summary
