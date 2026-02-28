"""
INTEGRATION — Page layout, header & footer.
Sets A4 page size, 2cm margins, portrait orientation,
and applies the header/footer template.
"""

from docx.shared import Mm, Pt, Twips, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from . import constants as C


def _rgb(hex_str: str) -> RGBColor:
    return RGBColor(int(hex_str[:2], 16), int(hex_str[2:4], 16), int(hex_str[4:], 16))


def apply_page_layout(doc):
    """Set A4 page size, 2cm margins, portrait orientation."""
    for section in doc.sections:
        section.page_width = Mm(C.PAGE_WIDTH_MM)
        section.page_height = Mm(C.PAGE_HEIGHT_MM)
        section.top_margin = Twips(C.MARGIN_DXA)
        section.bottom_margin = Twips(C.MARGIN_DXA)
        section.left_margin = Twips(C.MARGIN_DXA)
        section.right_margin = Twips(C.MARGIN_DXA)
        section.header_distance = Twips(C.HEADER_DIST_DXA)
        section.footer_distance = Twips(C.FOOTER_DIST_DXA)

        # Ensure portrait orientation
        sectPr = section._sectPr
        orient_el = sectPr.find(qn("w:pgSz"))
        if orient_el is not None:
            orient_el.set(qn("w:orient"), "portrait")


def _add_bottom_border_to_paragraph(paragraph, size: int, color: str):
    """Add a bottom border line to a paragraph (for header)."""
    pPr = paragraph._element.get_or_add_pPr()
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    border_xml = f"""
    <w:pBdr {nsdecls("w")}>
        <w:bottom w:val="single" w:sz="{size}" w:color="{color}" w:space="1"/>
    </w:pBdr>
    """
    pPr.append(parse_xml(border_xml))


def _add_top_border_to_paragraph(paragraph, size: int, color: str):
    """Add a top border line to a paragraph (for footer)."""
    pPr = paragraph._element.get_or_add_pPr()
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    border_xml = f"""
    <w:pBdr {nsdecls("w")}>
        <w:top w:val="single" w:sz="{size}" w:color="{color}" w:space="1"/>
    </w:pBdr>
    """
    pPr.append(parse_xml(border_xml))


def _add_right_tab_stop(paragraph, position_dxa: int):
    """Add a right-aligned tab stop to a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = paragraph._element.makeelement(qn("w:tabs"), {})
        pPr.append(tabs)
    tab = paragraph._element.makeelement(qn("w:tab"), {
        qn("w:val"): "right",
        qn("w:pos"): str(position_dxa),
    })
    tabs.append(tab)


def apply_header(doc, report_name: str = "Report"):
    """
    Apply the INTEGRATION header to all sections:
    Right-aligned, Arial 9pt, grey, with green bottom border.
    """
    header_text = C.HEADER_TEMPLATE.format(report_name=report_name)

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        # Clear existing
        for para in header.paragraphs:
            para.clear()

        if header.paragraphs:
            para = header.paragraphs[0]
        else:
            para = header.add_paragraph()

        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(header_text)
        run.font.name = C.FONT_FAMILY
        run.font.size = Pt(9)
        run.font.bold = False
        run.font.color.rgb = _rgb(C.HEADER_GREY)

        # Set rFonts
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = run._element.makeelement(qn("w:rFonts"), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), C.FONT_FAMILY)
        rFonts.set(qn("w:hAnsi"), C.FONT_FAMILY)

        # Bottom border
        _add_bottom_border_to_paragraph(para, C.HEADER_FOOTER_BORDER_SZ, C.GREEN_HAZE)

    return True


def apply_footer(doc, year: str = "2026", project_number: str = "PRJ-001"):
    """
    Apply the INTEGRATION footer to all sections:
    Left: copyright | project number, Right: bold green page number.
    Green top border.
    """
    footer_left_text = C.FOOTER_LEFT_TEMPLATE.format(year=year, project_number=project_number)

    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        # Clear existing
        for para in footer.paragraphs:
            para.clear()

        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()

        # Add right tab stop
        _add_right_tab_stop(para, C.FOOTER_TAB_STOP_DXA)

        # Top border
        _add_top_border_to_paragraph(para, C.HEADER_FOOTER_BORDER_SZ, C.GREEN_HAZE)

        # Left side text
        run_left = para.add_run(footer_left_text)
        run_left.font.name = C.FONT_FAMILY
        run_left.font.size = Pt(9)
        run_left.font.color.rgb = _rgb(C.HEADER_GREY)

        # Tab
        tab_run = para.add_run("\t")

        # Page number field
        run_page = para.add_run()
        run_page.font.name = C.FONT_FAMILY
        run_page.font.size = Pt(9)
        run_page.font.bold = True
        run_page.font.color.rgb = _rgb(C.GREEN_HAZE)

        # Insert PAGE field code
        fldChar_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_page._element.append(fldChar_begin)

        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_page_instr = para.add_run()
        run_page_instr._element.append(instrText)

        fldChar_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_page_end = para.add_run()
        run_page_end._element.append(fldChar_end)

    return True
