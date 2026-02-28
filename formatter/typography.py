"""
INTEGRATION — Typography enforcement.
Sets Arial font, correct size / bold / italic / colour on all runs.
"""

from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from . import constants as C


def _rgb(hex_str: str) -> RGBColor:
    """Convert a 6-char hex string to an RGBColor."""
    return RGBColor(int(hex_str[:2], 16), int(hex_str[2:4], 16), int(hex_str[4:], 16))


def apply_run_style(run, element_type: str):
    """Apply typography spec from constants.TYPO to a single run."""
    spec = C.TYPO.get(element_type)
    if not spec:
        return
    size_pt, bold, italic, colour_hex = spec
    run.font.name = C.FONT_FAMILY
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = _rgb(colour_hex)

    # Ensure Arial is set for East Asian / Complex Script fallback
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), C.FONT_FAMILY)
    rFonts.set(qn("w:hAnsi"), C.FONT_FAMILY)
    rFonts.set(qn("w:cs"), C.FONT_FAMILY)
    rFonts.set(qn("w:eastAsia"), C.FONT_FAMILY)


def apply_paragraph_style(paragraph, element_type: str):
    """Apply the typography spec to every run in a paragraph."""
    for run in paragraph.runs:
        apply_run_style(run, element_type)


def enforce_arial_everywhere(doc):
    """Walk ALL runs in the document and force Arial font family."""
    changes = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name and run.font.name != C.FONT_FAMILY:
                changes.append(run.font.name)
            run.font.name = C.FONT_FAMILY
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = run._element.makeelement(qn("w:rFonts"), {})
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:ascii"), C.FONT_FAMILY)
            rFonts.set(qn("w:hAnsi"), C.FONT_FAMILY)
            rFonts.set(qn("w:cs"), C.FONT_FAMILY)
            rFonts.set(qn("w:eastAsia"), C.FONT_FAMILY)

    # Also fix table cell runs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.name and run.font.name != C.FONT_FAMILY:
                            changes.append(run.font.name)
                        run.font.name = C.FONT_FAMILY
                        rPr = run._element.get_or_add_rPr()
                        rFonts = rPr.find(qn("w:rFonts"))
                        if rFonts is None:
                            rFonts = run._element.makeelement(qn("w:rFonts"), {})
                            rPr.insert(0, rFonts)
                        rFonts.set(qn("w:ascii"), C.FONT_FAMILY)
                        rFonts.set(qn("w:hAnsi"), C.FONT_FAMILY)
                        rFonts.set(qn("w:cs"), C.FONT_FAMILY)
                        rFonts.set(qn("w:eastAsia"), C.FONT_FAMILY)
    return changes
