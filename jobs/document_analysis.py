"""
AGENT ZEE — Documentation Analysis Engine
Analyzes .docx documents and produces comprehensive analysis reports with
readability metrics, structure breakdown, and quality assessment.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re
import math

GREEN = RGBColor(0x00, 0x99, 0x59)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x60, 0x60, 0x60)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
AMBER = RGBColor(0xFF, 0x8C, 0x00)
RED = RGBColor(0xCC, 0x00, 0x00)


def _set_run_style(run, size_pt, bold=False, italic=False, color=BLACK, font="Arial"):
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def _add_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    for r in para.runs:
        _set_run_style(r, 14 if level == 1 else 12 if level == 2 else 11,
                       bold=True, color=GREEN)
    return para


def _add_metric_row(table, label, value, rating=""):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = str(value)
    if len(row.cells) > 2:
        row.cells[2].text = rating


def _shade_header(row):
    for cell in row.cells:
        tc_pr = cell._element.get_or_add_tcPr()
        shd = tc_pr.makeelement(qn("w:shd"), {qn("w:fill"): "009959", qn("w:val"): "clear"})
        tc_pr.append(shd)
        for run in cell.paragraphs[0].runs:
            _set_run_style(run, 10, bold=True, color=WHITE)


def _analyze_structure(source_doc):
    """Analyze the structural elements of the document."""
    headings = {"level_1": [], "level_2": [], "level_3": [], "level_4": []}
    tables_count = len(source_doc.tables)
    images_count = 0
    lists_count = 0

    for para in source_doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        if style_name == "Heading 1":
            headings["level_1"].append(text)
        elif style_name == "Heading 2":
            headings["level_2"].append(text)
        elif style_name == "Heading 3":
            headings["level_3"].append(text)
        elif style_name == "Heading 4":
            headings["level_4"].append(text)

        if style_name.startswith("List"):
            lists_count += 1

    # Count images (inline shapes)
    for rel in source_doc.part.rels.values():
        if "image" in rel.reltype:
            images_count += 1

    return {
        "headings": headings,
        "total_headings": sum(len(v) for v in headings.values()),
        "tables": tables_count,
        "images": images_count,
        "lists": lists_count,
    }


def _analyze_content(source_doc):
    """Analyze the text content for metrics."""
    all_text = []
    paragraphs_count = 0

    for para in source_doc.paragraphs:
        text = para.text.strip()
        if text:
            all_text.append(text)
            paragraphs_count += 1

    full_text = " ".join(all_text)
    words = full_text.split()
    word_count = len(words)

    # Sentence count (approximate)
    sentences = re.split(r'[.!?]+', full_text)
    sentence_count = len([s for s in sentences if s.strip()])

    # Average words per sentence
    avg_words = round(word_count / max(sentence_count, 1), 1)

    # Syllable estimation (simple heuristic)
    def syllable_count(word):
        word = word.lower()
        vowels = "aeiou"
        count = sum(1 for i, c in enumerate(word) if c in vowels and (i == 0 or word[i-1] not in vowels))
        return max(count, 1)

    total_syllables = sum(syllable_count(w) for w in words)
    avg_syllables = round(total_syllables / max(word_count, 1), 2)

    # Flesch Reading Ease (approximate)
    if sentence_count > 0 and word_count > 0:
        flesch = 206.835 - 1.015 * (word_count / sentence_count) - 84.6 * (total_syllables / word_count)
        flesch = round(max(0, min(100, flesch)), 1)
    else:
        flesch = 0

    # Unique words
    unique_words = len(set(w.lower() for w in words if len(w) > 2))

    # Character count
    char_count = len(full_text)

    return {
        "word_count": word_count,
        "sentence_count": sentence_count,
        "paragraph_count": paragraphs_count,
        "char_count": char_count,
        "avg_words_per_sentence": avg_words,
        "avg_syllables_per_word": avg_syllables,
        "flesch_score": flesch,
        "unique_words": unique_words,
        "vocabulary_richness": round(unique_words / max(word_count, 1) * 100, 1),
    }


def _assess_quality(structure, content):
    """Generate a quality assessment with scores."""
    issues = []
    score = 100

    # Structure checks
    if structure["total_headings"] == 0:
        issues.append(("No headings detected — document lacks structure", -20))
    elif structure["total_headings"] < 3:
        issues.append(("Very few headings — document may need better organization", -10))

    if len(structure["headings"]["level_1"]) == 0:
        issues.append(("No top-level headings (H1) detected", -10))

    if structure["tables"] == 0:
        issues.append(("No tables found — consider adding data tables for clarity", -5))

    # Content checks
    if content["word_count"] < 500:
        issues.append(("Document is very short (< 500 words)", -15))
    elif content["word_count"] < 1000:
        issues.append(("Document is relatively short (< 1,000 words)", -5))

    if content["avg_words_per_sentence"] > 30:
        issues.append(("Sentences are too long on average (> 30 words)", -10))
    elif content["avg_words_per_sentence"] > 25:
        issues.append(("Sentences are slightly long on average (> 25 words)", -5))

    if content["flesch_score"] < 30:
        issues.append(("Readability is very low — text is difficult to read", -10))
    elif content["flesch_score"] < 50:
        issues.append(("Readability could be improved", -5))

    if content["vocabulary_richness"] < 20:
        issues.append(("Low vocabulary diversity — text may be repetitive", -5))

    for label, penalty in issues:
        score += penalty  # penalty is negative

    score = max(0, min(100, score))

    return {
        "score": score,
        "rating": "Excellent" if score >= 85 else "Good" if score >= 70 else "Needs Improvement" if score >= 50 else "Poor",
        "issues": issues,
    }


def _flesch_label(score):
    if score >= 80:
        return "Easy"
    elif score >= 60:
        return "Standard"
    elif score >= 40:
        return "Somewhat Difficult"
    elif score >= 20:
        return "Difficult"
    return "Very Difficult"


def process_analysis(input_path, output_path, **kwargs):
    """
    Analyze a document and produce a comprehensive analysis report.
    """
    analysis_depth = kwargs.get("analysis_depth", "Standard")

    source_doc = Document(input_path)
    structure = _analyze_structure(source_doc)
    content = _analyze_content(source_doc)
    quality = _assess_quality(structure, content)

    # ── Build output document ──
    doc = Document()

    # Title
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Document Analysis Report")
    _set_run_style(run, 20, bold=True, color=GREEN)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"Analysis Depth: {analysis_depth}")
    _set_run_style(run, 12, color=GREY)

    doc.add_paragraph()

    # ── Quality Score Overview ──
    _add_heading(doc, "1. Quality Score Overview", level=1)
    para = doc.add_paragraph()
    run = para.add_run(f"Overall Quality Score: {quality['score']}/100 — {quality['rating']}")
    _set_run_style(run, 14, bold=True,
                   color=GREEN if quality['score'] >= 70 else AMBER if quality['score'] >= 50 else RED)

    if quality["issues"]:
        _add_heading(doc, "Issues Identified", level=2)
        for issue_text, penalty in quality["issues"]:
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(f"{issue_text} ({penalty} points)")
            _set_run_style(run, 11)

    # ── Content Metrics ──
    _add_heading(doc, "2. Content Metrics", level=1)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Metric", "Value", "Assessment"]
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h
    _shade_header(tbl.rows[0])

    _add_metric_row(tbl, "Total Words", f"{content['word_count']:,}",
                    "✅ Good" if content['word_count'] >= 1000 else "⚠️ Short")
    _add_metric_row(tbl, "Sentences", f"{content['sentence_count']:,}", "")
    _add_metric_row(tbl, "Paragraphs", f"{content['paragraph_count']:,}", "")
    _add_metric_row(tbl, "Characters", f"{content['char_count']:,}", "")
    _add_metric_row(tbl, "Unique Words", f"{content['unique_words']:,}", "")
    _add_metric_row(tbl, "Vocabulary Richness", f"{content['vocabulary_richness']}%",
                    "✅" if content['vocabulary_richness'] >= 30 else "⚠️ Low diversity")
    _add_metric_row(tbl, "Avg Words/Sentence", f"{content['avg_words_per_sentence']}",
                    "✅" if content['avg_words_per_sentence'] <= 25 else "⚠️ Long sentences")
    _add_metric_row(tbl, "Flesch Reading Ease", f"{content['flesch_score']} ({_flesch_label(content['flesch_score'])})",
                    "✅" if content['flesch_score'] >= 50 else "⚠️ Difficult")

    # ── Structure Analysis ──
    _add_heading(doc, "3. Document Structure", level=1)

    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Element", "Count"]):
        tbl2.rows[0].cells[i].text = h
    _shade_header(tbl2.rows[0])

    _add_metric_row(tbl2, "Heading 1 (H1)", str(len(structure["headings"]["level_1"])))
    _add_metric_row(tbl2, "Heading 2 (H2)", str(len(structure["headings"]["level_2"])))
    _add_metric_row(tbl2, "Heading 3 (H3)", str(len(structure["headings"]["level_3"])))
    _add_metric_row(tbl2, "Heading 4 (H4)", str(len(structure["headings"]["level_4"])))
    _add_metric_row(tbl2, "Tables", str(structure["tables"]))
    _add_metric_row(tbl2, "Images/Figures", str(structure["images"]))
    _add_metric_row(tbl2, "List Items", str(structure["lists"]))

    # Heading hierarchy
    if structure["headings"]["level_1"]:
        _add_heading(doc, "Section Outline", level=2)
        for h1 in structure["headings"]["level_1"]:
            para = doc.add_paragraph()
            run = para.add_run(f"▸ {h1}")
            _set_run_style(run, 11, bold=True)

    # ── Deep Dive (if requested) ──
    if analysis_depth == "Deep Dive":
        doc.add_page_break()
        _add_heading(doc, "4. Detailed Content Analysis", level=1)

        # Top sections by content
        sections = {}
        current = "Preamble"
        sections[current] = 0
        for para in source_doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue
            if style_name.startswith("Heading"):
                current = text
                sections[current] = 0
            else:
                sections[current] += len(text.split())

        if sections:
            _add_heading(doc, "Word Distribution by Section", level=2)
            tbl3 = doc.add_table(rows=1, cols=3)
            tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, h in enumerate(["Section", "Word Count", "% of Total"]):
                tbl3.rows[0].cells[i].text = h
            _shade_header(tbl3.rows[0])

            total = max(sum(sections.values()), 1)
            for sec_name, wc in sorted(sections.items(), key=lambda x: -x[1])[:15]:
                _add_metric_row(tbl3, sec_name[:60], str(wc), f"{100*wc//total}%")

    # ── Recommendations ──
    doc.add_page_break()
    _add_heading(doc, "5. Recommendations", level=1)

    recs = _generate_recommendations(structure, content, quality)
    for idx, rec in enumerate(recs, 1):
        para = doc.add_paragraph()
        run = para.add_run(f"{idx}. {rec}")
        _set_run_style(run, 11)
        para.paragraph_format.space_after = Twips(120)

    # Save
    doc.save(output_path)

    summary = (
        f"DOCUMENTATION ANALYSIS SUMMARY\n"
        f"{'=' * 40}\n"
        f"Analysis Depth:        {analysis_depth}\n"
        f"Quality Score:         {quality['score']}/100 ({quality['rating']})\n"
        f"Word Count:            {content['word_count']:,}\n"
        f"Sentences:             {content['sentence_count']:,}\n"
        f"Flesch Reading Ease:   {content['flesch_score']} ({_flesch_label(content['flesch_score'])})\n"
        f"Headings Found:        {structure['total_headings']}\n"
        f"Tables:                {structure['tables']}\n"
        f"Issues Found:          {len(quality['issues'])}\n"
        f"Recommendations:       {len(recs)}\n"
        f"{'=' * 40}\n"
    )

    return summary


def _generate_recommendations(structure, content, quality):
    """Generate actionable recommendations based on analysis."""
    recs = []

    if structure["total_headings"] < 3:
        recs.append("Add more section headings to improve document navigation and readability.")

    if content["avg_words_per_sentence"] > 25:
        recs.append("Break down long sentences into shorter, clearer statements for better readability.")

    if content["flesch_score"] < 50:
        recs.append("Simplify language to improve readability — use shorter words and common terminology.")

    if content["vocabulary_richness"] < 25:
        recs.append("Diversify vocabulary to reduce repetitiveness and improve engagement.")

    if structure["tables"] == 0:
        recs.append("Consider adding tables to present data more clearly.")

    if structure["images"] == 0:
        recs.append("Consider adding figures, charts, or diagrams to visually support key points.")

    if content["word_count"] < 1000:
        recs.append("Expand the document content — current length may be insufficient for a comprehensive report.")

    if not recs:
        recs.append("Document is well-structured and meets quality standards. No major improvements needed.")

    return recs
