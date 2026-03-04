"""
AGENT ZEE — CV Reception & Rewriting Engine
Reformats CVs to match WB, EU, or ADB client template requirements.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Twips, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

GREEN = RGBColor(0x00,0x99,0x59); BLACK = RGBColor(0,0,0)
GREY = RGBColor(0x60,0x60,0x60); WHITE = RGBColor(0xFF,0xFF,0xFF)
BLUE = RGBColor(0x05,0x63,0xC1)

CV_TEMPLATES = {
    "World Bank (WB)": {
        "sections": [
            "Position Title and Key Qualifications",
            "Education",
            "Professional Certifications",
            "Employment Record (Starting with most recent)",
            "Languages",
            "Country Experience",
            "Publications and Training",
        ],
        "format_notes": [
            "WB format requires reverse chronological employment history",
            "Each assignment must show: dates, employer, position, description",
            "Country experience must list all countries worked in",
            "Maximum CV length: 5 pages per expert",
        ],
    },
    "European Union (EU)": {
        "sections": [
            "Proposed Position",
            "Personal Information",
            "Education and Training",
            "Language Skills (EU Scale: A1-C2)",
            "Professional Experience",
            "Other Relevant Information",
            "Declaration of Availability",
        ],
        "format_notes": [
            "EU format follows Europass CV structure",
            "Language skills must use CEFR scale (A1-C2)",
            "Professional experience in reverse chronological order",
            "Must include a declaration of availability and exclusivity",
        ],
    },
    "Asian Development Bank (ADB)": {
        "sections": [
            "Expert Name and Proposed Position",
            "Nationality and Date of Birth",
            "Education",
            "Professional Training",
            "Employment Record",
            "Detailed Description of Assignments",
            "Languages and Degree of Proficiency",
        ],
        "format_notes": [
            "ADB format requires detailed assignment descriptions",
            "Each assignment: client, location, dates, position, activities",
            "Must highlight experience in ADB member countries",
            "Include total years of professional experience",
        ],
    },
}

def _run(r,sz,bold=False,italic=False,color=BLACK):
    r.font.name="Arial"; r.font.size=Pt(sz)
    r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=color

def _heading(doc,text,lvl=1, level=None):
    if level is not None:
        lvl = level
    p=doc.add_heading(text,level=lvl)
    for r in p.runs: _run(r,14 if lvl==1 else 12 if lvl==2 else 11,bold=True,color=GREEN)

def _shade_hdr(row):
    for c in row.cells:
        tp=c._element.get_or_add_tcPr()
        tp.append(tp.makeelement(qn("w:shd"),{qn("w:fill"):"009959",qn("w:val"):"clear"}))
        for r in c.paragraphs[0].runs: _run(r,10,bold=True,color=WHITE)

def _extract_cv_data(src):
    """Extract CV information from source document."""
    data = {"name": "", "sections": {}, "all_text": []}
    cur_section = "Header"
    data["sections"][cur_section] = []

    for p in src.paragraphs:
        t = p.text.strip()
        if not t: continue
        sn = p.style.name if p.style else ""

        # Detect name (first bold or heading text)
        if not data["name"] and (sn.startswith("Heading") or (p.runs and p.runs[0].bold)):
            if len(t) < 60 and not any(k in t.lower() for k in ["education","experience","skill","language"]):
                data["name"] = t
                continue

        if sn.startswith("Heading") or (len(t) < 80 and t.isupper()) or \
           (p.runs and p.runs[0].bold and len(t) < 80):
            cur_section = t
            data["sections"][cur_section] = []
        else:
            data["sections"][cur_section].append(t)
            data["all_text"].append(t)

    return data

def _detect_experience_years(cv_data):
    """Try to detect years of experience from CV content."""
    all_text = " ".join(cv_data["all_text"])
    years = re.findall(r'(\d{4})', all_text)
    if years:
        years_int = [int(y) for y in years if 1970 <= int(y) <= 2026]
        if years_int:
            return 2026 - min(years_int)
    return 0

def _match_cv_sections(cv_data, template_sections):
    """Match extracted CV sections to template sections."""
    matched = {}
    for tmpl_sec in template_sections:
        best_match = None
        keywords = [w.lower() for w in tmpl_sec.split() if len(w) > 3]
        for cv_sec, content in cv_data["sections"].items():
            if any(k in cv_sec.lower() for k in keywords):
                best_match = (cv_sec, content)
                break
        matched[tmpl_sec] = best_match
    return matched


def process_cv_rewrite(input_path, output_path, **kwargs):
    client = kwargs.get("client_template", "World Bank (WB)")
    position = kwargs.get("position_title", "")
    min_years = int(kwargs.get("years_experience_required", "10"))

    src = Document(input_path)
    cv_data = _extract_cv_data(src)
    template = CV_TEMPLATES.get(client, CV_TEMPLATES["World Bank (WB)"])
    experience_years = _detect_experience_years(cv_data)

    doc = Document()

    # Header
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f"CURRICULUM VITAE"); _run(r,16,True,color=GREEN)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f"{client} Format"); _run(r,12,color=GREY)
    doc.add_paragraph()

    # Name and position
    if cv_data["name"]:
        p=doc.add_paragraph()
        r=p.add_run(f"Name: {cv_data['name']}"); _run(r,13,True)
    if position:
        p=doc.add_paragraph()
        r=p.add_run(f"Proposed Position: {position}"); _run(r,12,True,color=GREEN)
    doc.add_paragraph()

    # Build sections per template
    matched = _match_cv_sections(cv_data, template["sections"])
    populated = 0

    for idx, sec_name in enumerate(template["sections"], 1):
        _heading(doc, f"{idx}. {sec_name}", level=2)
        match = matched.get(sec_name)

        if match:
            populated += 1
            _, content = match
            for txt in content:
                p=doc.add_paragraph()
                r=p.add_run(txt); _run(r,11)
                p.paragraph_format.space_after=Twips(80)
        else:
            p=doc.add_paragraph()
            r=p.add_run(f"[To be populated — required for {client} submission]")
            _run(r,11,italic=True,color=GREY)
        doc.add_paragraph()

    # Compliance Notes
    doc.add_page_break()
    _heading(doc, "Compliance Notes", level=1)

    notes = []
    missing = [s for s in template["sections"] if not matched.get(s)]
    if missing:
        notes.append(f"⚠️ Missing sections ({len(missing)}): {', '.join(missing)}")
    if experience_years > 0 and experience_years < min_years:
        notes.append(f"⚠️ Detected {experience_years} years experience — minimum required is {min_years}")
    elif experience_years >= min_years:
        notes.append(f"✅ Experience: ~{experience_years} years (meets {min_years}-year minimum)")
    else:
        notes.append(f"⚠️ Could not detect years of experience — verify manually")
    notes.append(f"ℹ️ Populated {populated}/{len(template['sections'])} required sections")

    for note in notes:
        p=doc.add_paragraph(style="List Bullet")
        r=p.add_run(note); _run(r,11)

    # Format requirements
    _heading(doc, "Format Requirements", level=2)
    for note in template["format_notes"]:
        p=doc.add_paragraph(style="List Bullet")
        r=p.add_run(note); _run(r,10,color=GREY)

    doc.save(output_path)

    return (f"CV REWRITING SUMMARY\n{'='*40}\nCandidate: {cv_data['name'] or 'Unknown'}\n"
            f"Client Template: {client}\nPosition: {position or 'Not specified'}\n"
            f"Experience Detected: ~{experience_years} years\nMin Required: {min_years} years\n"
            f"Sections Populated: {populated}/{len(template['sections'])}\n"
            f"Missing Sections: {len(missing)}\n{'='*40}\n")
