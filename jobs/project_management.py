"""
INTEGRATION Agent — Project Management Engine
Generates WBS, Gantt tables, RACI matrix, and risk register from project docs.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re, math

GREEN = RGBColor(0x00,0x99,0x59); BLACK = RGBColor(0,0,0)
GREY = RGBColor(0x60,0x60,0x60); WHITE = RGBColor(0xFF,0xFF,0xFF)

def _run(r,sz,bold=False,italic=False,color=BLACK):
    r.font.name="Arial"; r.font.size=Pt(sz); r.font.bold=bold
    r.font.italic=italic; r.font.color.rgb=color

def _heading(doc,text,lvl=1):
    p=doc.add_heading(text,level=lvl)
    for r in p.runs: _run(r,14 if lvl==1 else 12,bold=True,color=GREEN)

def _shade_hdr(row):
    for c in row.cells:
        tp=c._element.get_or_add_tcPr()
        tp.append(tp.makeelement(qn("w:shd"),{qn("w:fill"):"009959",qn("w:val"):"clear"}))
        for r in c.paragraphs[0].runs: _run(r,10,bold=True,color=WHITE)

def _extract_activities(source_doc):
    acts, cur_phase, idx = [], "General", 0
    for p in source_doc.paragraphs:
        sn = p.style.name if p.style else ""
        t = p.text.strip()
        if not t: continue
        if sn.startswith("Heading"):
            cur_phase = t; continue
        if len(t) > 20:
            idx += 1
            acts.append({"id": f"A{idx:03d}", "phase": cur_phase, "desc": t[:120],
                         "duration": max(1, min(6, len(t.split())//15))})
    return acts

def process_project_management(input_path, output_path, **kwargs):
    project_name = kwargs.get("project_name", "Untitled Project")
    duration_months = int(kwargs.get("duration_months", "12"))

    src = Document(input_path)
    activities = _extract_activities(src)
    if not activities:
        activities = [{"id":"A001","phase":"Phase 1","desc":"Activity placeholder","duration":2}]

    doc = Document()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f"Project Management Report\n{project_name}"); _run(r,20,True,color=GREEN)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f"Duration: {duration_months} months"); _run(r,12,color=GREY)
    doc.add_paragraph()

    # WBS
    _heading(doc,"1. Work Breakdown Structure (WBS)")
    t=doc.add_table(rows=1,cols=4); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(["WBS #","Phase","Activity","Est. Duration (months)"]): t.rows[0].cells[i].text=h
    _shade_hdr(t.rows[0])
    phases_seen = {}
    for a in activities:
        rw=t.add_row()
        rw.cells[0].text=a["id"]; rw.cells[1].text=a["phase"][:40]
        rw.cells[2].text=a["desc"]; rw.cells[3].text=str(a["duration"])
        phases_seen[a["phase"]] = phases_seen.get(a["phase"], 0) + 1

    # Gantt
    doc.add_page_break()
    _heading(doc,"2. Gantt Timeline")
    months_cols = min(duration_months, 12)
    t2=doc.add_table(rows=1,cols=2+months_cols); t2.alignment=WD_TABLE_ALIGNMENT.CENTER
    t2.rows[0].cells[0].text="Activity"
    t2.rows[0].cells[1].text="Duration"
    for m in range(months_cols): t2.rows[0].cells[2+m].text=f"M{m+1}"
    _shade_hdr(t2.rows[0])

    running_month = 0
    for a in activities[:20]:
        rw=t2.add_row()
        rw.cells[0].text=a["desc"][:35]; rw.cells[1].text=f"{a['duration']}m"
        start = running_month % months_cols
        for m in range(start, min(start+a["duration"], months_cols)):
            tp=rw.cells[2+m]._element.get_or_add_tcPr()
            tp.append(tp.makeelement(qn("w:shd"),{qn("w:fill"):"009959",qn("w:val"):"clear"}))
            rw.cells[2+m].text="▓"
        running_month += a["duration"]

    # RACI
    doc.add_page_break()
    _heading(doc,"3. RACI Matrix")
    roles = ["Project Manager","Team Lead","Technical Expert","Client","Stakeholder"]
    t3=doc.add_table(rows=1,cols=1+len(roles)); t3.alignment=WD_TABLE_ALIGNMENT.CENTER
    t3.rows[0].cells[0].text="Activity"
    for i,role in enumerate(roles): t3.rows[0].cells[1+i].text=role
    _shade_hdr(t3.rows[0])
    raci_map = ["R","A","C","I","I"]
    for a in activities[:15]:
        rw=t3.add_row(); rw.cells[0].text=a["desc"][:35]
        for i,val in enumerate(raci_map): rw.cells[1+i].text=val

    # Risk Register
    doc.add_page_break()
    _heading(doc,"4. Risk Register")
    t4=doc.add_table(rows=1,cols=5); t4.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(["#","Risk Description","Likelihood","Impact","Mitigation"]):
        t4.rows[0].cells[i].text=h
    _shade_hdr(t4.rows[0])
    risks = [
        ("Scope creep due to changing requirements","Medium","High","Formal change control process"),
        ("Key personnel unavailability","Low","High","Cross-training and backup assignments"),
        ("Data quality issues","Medium","Medium","Data validation protocols"),
        ("Stakeholder engagement gaps","Medium","High","Regular communication schedule"),
        ("Budget overrun","Low","High","Monthly budget reviews and contingency fund"),
    ]
    for idx,(desc,lik,imp,mit) in enumerate(risks,1):
        rw=t4.add_row()
        rw.cells[0].text=str(idx); rw.cells[1].text=desc
        rw.cells[2].text=lik; rw.cells[3].text=imp; rw.cells[4].text=mit

    # Resource Allocation
    doc.add_page_break()
    _heading(doc,"5. Resource Allocation Matrix")
    t5=doc.add_table(rows=1,cols=4); t5.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(["Role","Allocation %","Person-Months","Key Responsibilities"]):
        t5.rows[0].cells[i].text=h
    _shade_hdr(t5.rows[0])
    allocs = [
        ("Project Manager","100%",str(duration_months),"Overall coordination and delivery"),
        ("Team Lead","80%",str(int(duration_months*0.8)),"Technical oversight and quality"),
        ("Technical Expert","60%",str(int(duration_months*0.6)),"Core deliverable production"),
        ("Support Staff","40%",str(int(duration_months*0.4)),"Administrative and logistics"),
    ]
    for role,pct,pm,resp in allocs:
        rw=t5.add_row()
        rw.cells[0].text=role; rw.cells[1].text=pct
        rw.cells[2].text=pm; rw.cells[3].text=resp

    doc.save(output_path)

    return (f"PROJECT MANAGEMENT SUMMARY\n{'='*40}\nProject: {project_name}\n"
            f"Duration: {duration_months} months\nActivities Extracted: {len(activities)}\n"
            f"Phases Identified: {len(phases_seen)}\nRisks Documented: {len(risks)}\n"
            f"Sections Generated: WBS, Gantt, RACI, Risk Register, Resource Allocation\n{'='*40}\n")
