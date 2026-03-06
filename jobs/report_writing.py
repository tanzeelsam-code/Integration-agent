"""
AGENT ZEE — Report Writing Engine
Transforms raw content into professionally structured reports for WB/EU/ADB.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Twips, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

GREEN = RGBColor(0x00,0x99,0x59); BLACK = RGBColor(0,0,0)
GREY = RGBColor(0x60,0x60,0x60); WHITE = RGBColor(0xFF,0xFF,0xFF)

REPORT_TEMPLATES = {
    "World Bank": ["Executive Summary","Introduction","Background and Context",
        "Project Description","Implementation Arrangements","Results Framework",
        "Monitoring and Evaluation","Financial Analysis","Annexes"],
    "European Union": ["Executive Summary","Context and Rationale","Project Design",
        "Methodology","Implementation Plan","Sustainability","Cross-cutting Issues",
        "Budget and Financial Plan","Annexes"],
    "ADB": ["Executive Summary","Introduction","Sector Assessment","Project Description",
        "Implementation","Financial Management","Risk Assessment","Appendixes"],
    "Generic Professional": ["Executive Summary","Introduction","Background",
        "Analysis","Findings","Recommendations","Conclusion","Appendices"],
}

AGENT_WORKSTREAMS = [
    {
        "name": "Agent 1 - Lead Writer",
        "focus": "Upgrade report drafting quality, structure, and coherence.",
        "deliverable": "Primary narrative draft for all sections.",
    },
    {
        "name": "Agent 2 - Technical Reviewer",
        "focus": "Review technical accuracy, evidence use, and alignment with source content.",
        "deliverable": "Review notes and flagged technical gaps.",
    },
    {
        "name": "Agent 3 - Revision Specialist",
        "focus": "Update reviewed reports with agreed improvements.",
        "deliverable": "Revised version with integrated reviewer changes.",
    },
    {
        "name": "Agent 4 - Comment Resolution Lead",
        "focus": "Address stakeholder comments and track final closure status.",
        "deliverable": "Comment response log with action status.",
    },
]


def _run(r,sz,bold=False,italic=False,color=BLACK):
    r.font.name="Arial"; r.font.size=Pt(sz)
    r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=color

def _heading(doc,text,lvl=1):
    p=doc.add_heading(text,level=lvl)
    for r in p.runs: _run(r,14 if lvl==1 else 12 if lvl==2 else 11,bold=True,color=GREEN)

def _shade_hdr(row):
    for c in row.cells:
        tp=c._element.get_or_add_tcPr()
        tp.append(tp.makeelement(qn("w:shd"),{qn("w:fill"):"009959",qn("w:val"):"clear"}))
        for r in c.paragraphs[0].runs: _run(r,10,bold=True,color=WHITE)

def _extract_content(src):
    secs, cur, txts = {}, "Unstructured Content", {}
    cur_texts = []
    for p in src.paragraphs:
        sn = p.style.name if p.style else ""
        t = p.text.strip()
        if not t: continue
        if sn.startswith("Heading"):
            if cur_texts: secs[cur] = cur_texts
            cur, cur_texts = t, []
        else:
            cur_texts.append(t)
    if cur_texts: secs[cur] = cur_texts
    return secs

def _generate_exec_summary(content):
    all_text = []
    for texts in content.values():
        all_text.extend(texts[:2])
    words = " ".join(all_text).split()
    if len(words) > 150:
        return " ".join(words[:150]) + "..."
    return " ".join(words) if words else "[Executive summary to be developed based on report content.]"


def _render_agent_plan(doc, assigned_agents):
    _heading(doc, "Delivery Workflow and Agent Allocation")
    p = doc.add_paragraph()
    r = p.add_run(
        "This report workflow uses a 4-agent team to ensure writing quality, formal review,"
        " revision of reviewed drafts, and closure of comments."
    )
    _run(r, 11)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    hdr.cells[0].text = "Agent"
    hdr.cells[1].text = "Primary Responsibility"
    hdr.cells[2].text = "Output"
    hdr.cells[3].text = "Status"
    _shade_hdr(hdr)

    active_agents = AGENT_WORKSTREAMS[:assigned_agents]
    for role in active_agents:
        row = table.add_row()
        row.cells[0].text = role["name"]
        row.cells[1].text = role["focus"]
        row.cells[2].text = role["deliverable"]
        row.cells[3].text = "Active"

    doc.add_paragraph()


def _add_review_capacity_sections(doc, style):
    _heading(doc, "Review and Quality Assurance")
    p = doc.add_paragraph()
    r = p.add_run(
        f"[Review checkpoint to be completed. Verify structure, evidence, consistency, and {style} compliance.]"
    )
    _run(r, 11, italic=True, color=GREY)
    for item in [
        "List sections reviewed and reviewer observations.",
        "Record critical and non-critical findings.",
        "Confirm decisions for each finding (accept, revise, defer).",
    ]:
        b = doc.add_paragraph(style="List Bullet")
        rb = b.add_run(item)
        _run(rb, 11, color=GREY)
    doc.add_paragraph()

    _heading(doc, "Updated Version of Reviewed Report")
    p = doc.add_paragraph()
    r = p.add_run(
        "[Revised text to be inserted after implementing accepted review findings.]"
    )
    _run(r, 11, italic=True, color=GREY)
    for item in [
        "Summarize major updates made after review.",
        "Identify sections that changed and why.",
        "Capture any unresolved review items.",
    ]:
        b = doc.add_paragraph(style="List Bullet")
        rb = b.add_run(item)
        _run(rb, 11, color=GREY)
    doc.add_paragraph()

    _heading(doc, "Comment Resolution and Response Log")
    p = doc.add_paragraph()
    r = p.add_run(
        "[Track comments, responses, actions taken, and closure status for final submission.]"
    )
    _run(r, 11, italic=True, color=GREY)

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    hdr.cells[0].text = "Comment ID"
    hdr.cells[1].text = "Comment"
    hdr.cells[2].text = "Response"
    hdr.cells[3].text = "Action Taken"
    hdr.cells[4].text = "Status"
    _shade_hdr(hdr)

    for cid in ["C-01", "C-02", "C-03"]:
        row = table.add_row()
        row.cells[0].text = cid
        row.cells[1].text = "[Enter comment]"
        row.cells[2].text = "[Enter response]"
        row.cells[3].text = "[Describe update]"
        row.cells[4].text = "Open"

    doc.add_paragraph()


def process_report(input_path, output_path, **kwargs):
    style = kwargs.get("report_style", "Generic Professional")
    title = kwargs.get("report_title", "Report")
    inc_exec = kwargs.get("include_exec_summary", "Yes")
    try:
        assigned_agents = int(kwargs.get("assigned_agents", "4"))
    except (TypeError, ValueError):
        assigned_agents = 4
    assigned_agents = 4 if assigned_agents < 4 else min(assigned_agents, 4)

    src = Document(input_path)
    content = _extract_content(src)
    total_words = sum(len(t.split()) for txts in content.values() for t in txts)
    template = REPORT_TEMPLATES.get(style, REPORT_TEMPLATES["Generic Professional"])

    doc = Document()

    # Cover
    for _ in range(3):
        doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(title or "Professional Report"); _run(r,24,True,color=GREEN)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f"\n{style} Format"); _run(r,14,color=GREY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("\n\nPrepared by AGENT ZEE"); _run(r,11,italic=True,color=GREY)
    doc.add_page_break()

    # TOC
    _heading(doc,"Table of Contents")
    for idx,sec in enumerate(template,1):
        p=doc.add_paragraph()
        r=p.add_run(f"{idx}. {sec}"); _run(r,11)
    p = doc.add_paragraph()
    r = p.add_run(f"{len(template) + 1}. Review and Quality Assurance")
    _run(r, 11)
    p = doc.add_paragraph()
    r = p.add_run(f"{len(template) + 2}. Updated Version of Reviewed Report")
    _run(r, 11)
    p = doc.add_paragraph()
    r = p.add_run(f"{len(template) + 3}. Comment Resolution and Response Log")
    _run(r, 11)
    doc.add_page_break()
    _render_agent_plan(doc, assigned_agents)

    # Sections
    populated = 0
    for idx,sec_name in enumerate(template,1):
        _heading(doc,f"{idx}. {sec_name}")

        # Exec summary special handling
        if "executive summary" in sec_name.lower() and inc_exec == "Yes":
            summary_text = _generate_exec_summary(content)
            p=doc.add_paragraph(); r=p.add_run(summary_text); _run(r,11)
            p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            populated += 1
            doc.add_paragraph()
            continue

        # Match source content
        matched = []
        for src_key, src_texts in content.items():
            if any(w in src_key.lower() for w in sec_name.lower().split() if len(w) > 3):
                matched.extend(src_texts)

        if matched:
            populated += 1
            for txt in matched:
                p=doc.add_paragraph(); r=p.add_run(txt); _run(r,11)
                p.paragraph_format.space_after=Twips(120)
                p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            p=doc.add_paragraph()
            r=p.add_run(f"[Section to be developed. Per {style} guidelines, this section should cover:]")
            _run(r,11,italic=True,color=GREY)
            for guide in _section_guidance(sec_name):
                p=doc.add_paragraph(style="List Bullet")
                r=p.add_run(guide); _run(r,11,color=GREY)
        doc.add_paragraph()

    _add_review_capacity_sections(doc, style)

    doc.save(output_path)

    return (f"REPORT WRITING SUMMARY\n{'='*40}\nReport Style: {style}\nTitle: {title}\n"
            f"Assigned Agents: {assigned_agents}\n"
            f"Source Words: {total_words:,}\nTemplate Sections: {len(template)}\n"
            f"Sections Populated: {populated}\nSections Need Dev: {len(template)-populated}\n"
            f"Review Capacity: Enabled\nRevision Capacity: Enabled\nComment Resolution Capacity: Enabled\n{'='*40}\n")

def _section_guidance(name):
    guides = {
        "introduction": ["State the purpose and scope","Define key terms","Outline document structure"],
        "background": ["Provide historical context","Summarize previous work","Identify key stakeholders"],
        "analysis": ["Present data and findings","Provide interpretations","Compare with benchmarks"],
        "findings": ["List key findings by theme","Support with evidence","Prioritize by significance"],
        "recommendations": ["Provide actionable recommendations","Link to findings","Include timelines"],
        "conclusion": ["Summarize main findings","Restate recommendations","Identify next steps"],
    }
    for key, bullets in guides.items():
        if key in name.lower(): return bullets
    return [f"Develop content per {name} requirements","Include supporting data","Reference source documents"]
