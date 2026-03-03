"""
INTEGRATION Agent — JIS Mapping Engine
Joint Implementation Support mapping — Results Framework, LogFrame, M&E matrix.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

GREEN = RGBColor(0x00,0x99,0x59); BLACK = RGBColor(0,0,0)
GREY = RGBColor(0x60,0x60,0x60); WHITE = RGBColor(0xFF,0xFF,0xFF)

def _run(r,sz,bold=False,italic=False,color=BLACK):
    r.font.name="Arial"; r.font.size=Pt(sz)
    r.font.bold=bold; r.font.italic=italic; r.font.color.rgb=color

def _heading(doc,text,lvl=1):
    p=doc.add_heading(text,level=lvl)
    for r in p.runs: _run(r,14 if lvl==1 else 12 if lvl==2 else 11,bold=True,color=GREEN)

def _shade_hdr(row):
    for c in row.cells:
        tp=c._element.get_or_add_tcPr()
        tp.append(tp.makeelement(qn("w:shd"),{qn("w:fill"):"009959",qn("w:val"):"clear"}))
        for r in c.paragraphs[0].runs: _run(r,9,bold=True,color=WHITE)

def _extract_objectives(src):
    objectives, activities, indicators = [], [], []
    keyword_obj = ["objective","goal","outcome","purpose","impact"]
    keyword_act = ["activity","task","action","deliverable","output"]
    keyword_ind = ["indicator","target","milestone","baseline","measure"]

    for p in src.paragraphs:
        t = p.text.strip().lower()
        if not t or len(t) < 10: continue
        if any(k in t for k in keyword_obj):
            objectives.append(p.text.strip())
        elif any(k in t for k in keyword_act):
            activities.append(p.text.strip())
        elif any(k in t for k in keyword_ind):
            indicators.append(p.text.strip())

    # Fallback: use headings as objectives, body as activities
    if not objectives:
        for p in src.paragraphs:
            sn = p.style.name if p.style else ""
            if sn.startswith("Heading") and p.text.strip():
                objectives.append(p.text.strip())
    if not activities:
        for p in src.paragraphs:
            sn = p.style.name if p.style else ""
            if not sn.startswith("Heading") and len(p.text.strip()) > 30:
                activities.append(p.text.strip()[:120])
                if len(activities) >= 10: break

    return objectives[:10], activities[:15], indicators[:10]

def _build_logframe(doc, objectives, activities, indicators):
    _heading(doc, "Logical Framework (LogFrame)")
    headers = ["Level","Description","Indicators","Means of Verification","Assumptions"]
    t = doc.add_table(rows=1,cols=5); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers): t.rows[0].cells[i].text=h
    _shade_hdr(t.rows[0])

    # Goal level
    rw=t.add_row()
    rw.cells[0].text="Goal"
    rw.cells[1].text=objectives[0] if objectives else "[To be defined]"
    rw.cells[2].text=indicators[0] if indicators else "[TBD]"
    rw.cells[3].text="Project reports, M&E data"
    rw.cells[4].text="Enabling policy environment"

    # Purpose
    rw=t.add_row()
    rw.cells[0].text="Purpose"
    rw.cells[1].text=objectives[1] if len(objectives)>1 else "[To be defined]"
    rw.cells[2].text=indicators[1] if len(indicators)>1 else "[TBD]"
    rw.cells[3].text="Progress reports, surveys"
    rw.cells[4].text="Stakeholder commitment"

    # Outputs
    for i,act in enumerate(activities[:5]):
        rw=t.add_row()
        rw.cells[0].text=f"Output {i+1}"
        rw.cells[1].text=act[:100]
        rw.cells[2].text=indicators[i+2] if len(indicators)>i+2 else f"Output {i+1} delivered"
        rw.cells[3].text="Deliverable review, field visits"
        rw.cells[4].text="Resources available on time"

    # Activities
    for i,act in enumerate(activities[5:10]):
        rw=t.add_row()
        rw.cells[0].text=f"Activity {i+1}"
        rw.cells[1].text=act[:100]
        rw.cells[2].text="Completion report"
        rw.cells[3].text="Activity logs"
        rw.cells[4].text="No major disruptions"

def _build_results_framework(doc, objectives, activities, indicators):
    _heading(doc, "Results Framework")
    headers = ["Result Level","Description","Indicator","Baseline","Target","Data Source","Frequency"]
    t = doc.add_table(rows=1,cols=7); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers): t.rows[0].cells[i].text=h
    _shade_hdr(t.rows[0])

    levels = [("PDO",objectives[:2]),("Intermediate Outcome",objectives[2:5]),
              ("Output",activities[:5])]
    for level_name, items in levels:
        for i,item in enumerate(items):
            rw=t.add_row()
            rw.cells[0].text=level_name
            rw.cells[1].text=item[:80]
            rw.cells[2].text=indicators[i] if i < len(indicators) else "[TBD]"
            rw.cells[3].text="[Baseline]"; rw.cells[4].text="[Target]"
            rw.cells[5].text="Project M&E"; rw.cells[6].text="Annual"

def _build_me_matrix(doc, objectives, activities, indicators):
    _heading(doc, "M&E Mapping Matrix")
    headers = ["Indicator","Type","Frequency","Responsible","Data Source","Reporting"]
    t = doc.add_table(rows=1,cols=6); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers): t.rows[0].cells[i].text=h
    _shade_hdr(t.rows[0])

    all_ind = indicators if indicators else [f"Indicator {i+1}" for i in range(5)]
    types = ["Outcome","Output","Process"]
    freqs = ["Annual","Semi-annual","Quarterly"]
    for i,ind in enumerate(all_ind[:10]):
        rw=t.add_row()
        rw.cells[0].text=ind[:80]
        rw.cells[1].text=types[i%3]
        rw.cells[2].text=freqs[i%3]
        rw.cells[3].text="M&E Officer"
        rw.cells[4].text="Project records"
        rw.cells[5].text="Progress reports"

def process_jis_mapping(input_path, output_path, **kwargs):
    framework_type = kwargs.get("framework_type", "Full Package")
    sector = kwargs.get("sector", "Energy")

    src = Document(input_path)
    objectives, activities, indicators = _extract_objectives(src)

    doc = Document()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("JIS Mapping Report"); _run(r,20,True,color=GREEN)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f"Sector: {sector} | Framework: {framework_type}"); _run(r,12,color=GREY)
    doc.add_paragraph()

    # Extraction summary
    _heading(doc, "1. Extraction Summary")
    p=doc.add_paragraph()
    r=p.add_run(f"Objectives identified: {len(objectives)}\n"
                f"Activities identified: {len(activities)}\n"
                f"Indicators identified: {len(indicators)}"); _run(r,11)
    doc.add_paragraph()

    sections_built = []
    if framework_type in ("Logical Framework (LogFrame)", "Full Package"):
        doc.add_page_break()
        _heading(doc, "2. Logical Framework")
        _build_logframe(doc, objectives, activities, indicators)
        sections_built.append("LogFrame")

    if framework_type in ("Results Framework", "Full Package"):
        doc.add_page_break()
        _heading(doc, "3. Results Framework")
        _build_results_framework(doc, objectives, activities, indicators)
        sections_built.append("Results Framework")

    if framework_type in ("M&E Matrix", "Full Package"):
        doc.add_page_break()
        _heading(doc, "4. M&E Mapping Matrix")
        _build_me_matrix(doc, objectives, activities, indicators)
        sections_built.append("M&E Matrix")

    doc.save(output_path)

    return (f"JIS MAPPING SUMMARY\n{'='*40}\nSector: {sector}\nFramework Type: {framework_type}\n"
            f"Objectives Found: {len(objectives)}\nActivities Found: {len(activities)}\n"
            f"Indicators Found: {len(indicators)}\nSections Generated: {', '.join(sections_built)}\n{'='*40}\n")
