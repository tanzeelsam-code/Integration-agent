"""
INTEGRATION Agent — Document Comparison Engine
Compares two .docx documents with structural and content gap analysis.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import difflib

GREEN = RGBColor(0x00, 0x99, 0x59)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x60, 0x60, 0x60)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _run(r, sz, bold=False, italic=False, color=BLACK):
    r.font.name = "Arial"; r.font.size = Pt(sz)
    r.font.bold = bold; r.font.italic = italic; r.font.color.rgb = color

def _heading(doc, text, lvl=1):
    p = doc.add_heading(text, level=lvl)
    for r in p.runs: _run(r, 14 if lvl==1 else 12, bold=True, color=GREEN)

def _shade_hdr(row):
    for c in row.cells:
        tp = c._element.get_or_add_tcPr()
        tp.append(tp.makeelement(qn("w:shd"), {qn("w:fill"): "009959", qn("w:val"): "clear"}))
        for r in c.paragraphs[0].runs: _run(r, 10, bold=True, color=WHITE)

def _shade(cell, fill):
    tp = cell._element.get_or_add_tcPr()
    tp.append(tp.makeelement(qn("w:shd"), {qn("w:fill"): fill, qn("w:val"): "clear"}))

def _extract(doc):
    secs, cur_h, cur_c = [], None, []
    for p in doc.paragraphs:
        sn = p.style.name if p.style else ""
        t = p.text.strip()
        if sn.startswith("Heading") and t:
            if cur_h is not None:
                secs.append({"heading": cur_h, "content": " ".join(cur_c), "wc": len(" ".join(cur_c).split())})
            cur_h, cur_c = t, []
        elif t:
            cur_c.append(t)
    if cur_h is not None:
        secs.append({"heading": cur_h, "content": " ".join(cur_c), "wc": len(" ".join(cur_c).split())})
    elif cur_c:
        secs.append({"heading": "(No headings)", "content": " ".join(cur_c), "wc": len(" ".join(cur_c).split())})
    return {"sections": secs, "words": sum(s["wc"] for s in secs),
            "tables": len(doc.tables), "headings": [s["heading"] for s in secs]}

def _sim(a, b):
    if not a and not b: return 1.0
    if not a or not b: return 0.0
    return difflib.SequenceMatcher(None, a.lower(), b.lower()).ratio()

def _match(da, db):
    matched, used = [], set()
    for sa in da["sections"]:
        best, bs = None, 0.0
        for i, sb in enumerate(db["sections"]):
            if i in used: continue
            s = _sim(sa["heading"], sb["heading"])
            if s > bs: bs, best = s, (i, sb)
        if best and bs >= 0.3:
            used.add(best[0])
            cs = _sim(sa["content"], best[1]["content"])
            matched.append((sa, best[1], round(bs*100,1), round(cs*100,1)))
        else:
            matched.append((sa, None, 0, 0))
    for i, sb in enumerate(db["sections"]):
        if i not in used: matched.append((None, sb, 0, 0))
    return matched


def process_comparison(input_path, output_path, **kwargs):
    mode = kwargs.get("comparison_mode", "Full Comparison")
    if isinstance(input_path, (list, tuple)) and len(input_path) >= 2:
        pa, pb = input_path[0], input_path[1]
    else:
        pa, pb = input_path, kwargs.get("second_file", input_path)

    da, db = _extract(Document(pa)), _extract(Document(pb))
    all_a = " ".join(s["content"] for s in da["sections"])
    all_b = " ".join(s["content"] for s in db["sections"])
    overall = round(_sim(all_a, all_b)*100, 1)

    doc = Document()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Document Comparison Report"); _run(r, 20, True, color=GREEN)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Mode: {mode}"); _run(r, 12, color=GREY)

    _heading(doc, "1. Overview Comparison")
    t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Metric","Document A","Document B"]): t.rows[0].cells[i].text = h
    _shade_hdr(t.rows[0])
    for lbl, va, vb in [("Words",f"{da['words']:,}",f"{db['words']:,}"),
                         ("Sections",str(len(da['sections'])),str(len(db['sections']))),
                         ("Tables",str(da['tables']),str(db['tables']))]:
        rw = t.add_row(); rw.cells[0].text=lbl; rw.cells[1].text=va; rw.cells[2].text=vb
    p = doc.add_paragraph(); r = p.add_run(f"\nOverall Similarity: {overall}%")
    _run(r, 13, True, color=GREEN)

    matched = _match(da, db)
    if mode in ("Full Comparison","Structure Only"):
        doc.add_page_break(); _heading(doc, "2. Structural Comparison")
        t2 = doc.add_table(rows=1, cols=5); t2.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(["#","Doc A Section","Doc B Section","Head %","Content %"]):
            t2.rows[0].cells[i].text = h
        _shade_hdr(t2.rows[0])
        for idx,(sa,sb,hs,cs) in enumerate(matched,1):
            rw = t2.add_row(); rw.cells[0].text=str(idx)
            rw.cells[1].text = (sa["heading"] if sa else "— Missing —")[:50]
            rw.cells[2].text = (sb["heading"] if sb else "— Missing —")[:50]
            rw.cells[3].text = f"{hs}%"; rw.cells[4].text = f"{cs}%"
            if not sa: _shade(rw.cells[1], "FFC7CE")
            if not sb: _shade(rw.cells[2], "FFC7CE")
            _shade(rw.cells[4], "C6EFCE" if cs>=70 else "FFEB9C" if cs>=30 else "FFC7CE")

    only_a = [h for h in da["headings"] if h.lower() not in {x.lower() for x in db["headings"]}]
    only_b = [h for h in db["headings"] if h.lower() not in {x.lower() for x in da["headings"]}]
    if mode in ("Full Comparison","Content Only"):
        doc.add_page_break(); _heading(doc, "3. Gap Analysis")
        _heading(doc, "Only in Document A", 2)
        for h in (only_a or ["None"]):
            p = doc.add_paragraph(style="List Bullet"); r = p.add_run(h); _run(r, 11)
        _heading(doc, "Only in Document B", 2)
        for h in (only_b or ["None"]):
            p = doc.add_paragraph(style="List Bullet"); r = p.add_run(h); _run(r, 11)

    doc.save(output_path)
    sm = sum(1 for _,_,_,cs in matched if cs>=50)
    return (f"DOCUMENT COMPARISON SUMMARY\n{'='*40}\nMode: {mode}\nSimilarity: {overall}%\n"
            f"Doc A: {da['words']:,} words, {len(da['sections'])} sections\n"
            f"Doc B: {db['words']:,} words, {len(db['sections'])} sections\n"
            f"Strong Matches: {sm}\nUnique to A: {len(only_a)}\nUnique to B: {len(only_b)}\n{'='*40}\n")
