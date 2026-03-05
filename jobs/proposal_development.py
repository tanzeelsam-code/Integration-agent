"""
AGENT ZEE — Proposal Development Engine
Analyzes draft proposals/TORs and generates structured proposal documents
compliant with donor/client requirements.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re


# ─── CLIENT TEMPLATES ────────────────────────────────────────────────
CLIENT_SECTIONS = {
    "World Bank (WB)": [
        "Executive Summary",
        "Understanding of the Assignment",
        "Technical Approach and Methodology",
        "Work Plan and Schedule",
        "Team Composition and Task Assignment",
        "Qualifications and Experience of Key Staff",
        "Past Performance and References",
    ],
    "European Union (EU)": [
        "Executive Summary",
        "Relevance of the Action",
        "Methodology and Approach",
        "Sustainability and Impact",
        "Work Plan and Timeline",
        "Budget Justification",
        "Consortium / Team Structure",
        "Logical Framework Matrix",
    ],
    "Asian Development Bank (ADB)": [
        "Introduction and Context",
        "Understanding of the TOR",
        "Technical Approach",
        "Methodology",
        "Work Plan",
        "Staffing Schedule",
        "Counterpart Facilities Required",
    ],
    "GIZ": [
        "Executive Summary",
        "Context and Problem Analysis",
        "Objectives and Expected Results",
        "Methodology and Implementation Concept",
        "Work Plan and Deliverables",
        "Team Composition and Roles",
        "Quality Assurance and Risk Management",
    ],
    "KfW": [
        "Executive Summary",
        "Project Background and Rationale",
        "Approach and Methodology",
        "Implementation Arrangements",
        "Schedule and Milestones",
        "Team Structure and Expert Inputs",
        "Monitoring and Reporting",
    ],
    "AFD": [
        "Executive Summary",
        "Context and Strategic Alignment",
        "Methodology and Technical Approach",
        "Results Framework and Indicators",
        "Implementation Plan",
        "Team Composition and Governance",
        "Sustainability and Risk Mitigation",
    ],
    "EBRD": [
        "Executive Summary",
        "Assignment Understanding",
        "Technical Methodology",
        "Work Plan and Timeline",
        "Team Composition and Key Experts",
        "Relevant Experience and References",
        "Project Management and Quality Control",
    ],
    "Other": [
        "Executive Summary",
        "Background and Context",
        "Technical Approach",
        "Methodology",
        "Work Plan",
        "Team Composition",
        "Budget Overview",
    ],
}

GREEN = RGBColor(0x00, 0x99, 0x59)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x60, 0x60, 0x60)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _set_run_style(run, size_pt, bold=False, italic=False, color=BLACK, font="Arial"):
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def _add_heading_styled(doc, text, level=1):
    """Add a heading with AGENT ZEE green styling."""
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        _set_run_style(run, 14 if level == 1 else 12 if level == 2 else 11,
                       bold=True, color=GREEN)
    return para


def _extract_source_content(source_doc):
    """
    Extract all text content from the source document, organized by
    detected sections.
    """
    sections = {}
    current_section = "General Content"
    sections[current_section] = []

    for para in source_doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading") or (
            len(text) < 120 and text.isupper()
        ) or (
            len(text) < 120 and para.runs and para.runs[0].bold
        ):
            current_section = text
            sections[current_section] = []
        else:
            sections[current_section].append(text)

    return sections


def _generate_compliance_checklist(client_type, source_sections):
    """Check which required sections are covered in the source document."""
    required = CLIENT_SECTIONS.get(client_type, CLIENT_SECTIONS["Other"])
    checklist = []

    source_keys_lower = [k.lower() for k in source_sections.keys()]

    for section in required:
        found = any(section.lower() in sk for sk in source_keys_lower)
        checklist.append({
            "section": section,
            "status": "✅ Covered" if found else "⚠️ Missing / Needs Development",
            "found": found,
        })

    return checklist


def process_proposal(input_path, output_path, **kwargs):
    """
    Process a draft proposal/TOR and generate a structured proposal document.

    Returns:
        Summary string describing what was generated.
    """
    client_type = kwargs.get("client_type", "World Bank (WB)")
    assigned_agents = kwargs.get("assigned_agents", "4")
    project_title = kwargs.get("project_title", "Untitled Project")
    country = kwargs.get("country", "Pakistan")

    # Read source document
    source_doc = Document(input_path)
    source_content = _extract_source_content(source_doc)
    source_word_count = sum(len(t.split()) for texts in source_content.values() for t in texts)

    # Create output document
    doc = Document()

    # ── Cover Page ──
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("\n\n\n")
    _set_run_style(run, 11)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(project_title or "Technical Proposal")
    _set_run_style(run, 22, bold=True, color=GREEN)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"\n{client_type} Compliant Submission")
    _set_run_style(run, 14, color=GREY)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"\nCountry/Region: {country}")
    _set_run_style(run, 12, color=GREY)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"\nAssigned Agents: {assigned_agents}")
    _set_run_style(run, 12, color=GREY)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("\n\nPrepared by AGENT ZEE")
    _set_run_style(run, 11, italic=True, color=GREY)

    doc.add_page_break()

    # ── Table of Contents placeholder ──
    _add_heading_styled(doc, "Table of Contents", level=1)
    para = doc.add_paragraph()
    run = para.add_run("[Auto-generated Table of Contents — update field after editing]")
    _set_run_style(run, 10, italic=True, color=GREY)
    doc.add_page_break()

    # ── Required Sections ──
    required_sections = CLIENT_SECTIONS.get(client_type, CLIENT_SECTIONS["Other"])
    sections_populated = 0

    for idx, section_name in enumerate(required_sections, 1):
        _add_heading_styled(doc, f"{idx}. {section_name}", level=1)

        # Try to find matching content from the source document
        matched_content = []
        for src_key, src_texts in source_content.items():
            if section_name.lower() in src_key.lower() or any(
                section_name.lower().split()[0] in src_key.lower().split()
                for _ in [1] if section_name.lower().split()
            ):
                matched_content.extend(src_texts)

        if matched_content:
            sections_populated += 1
            for text in matched_content:
                para = doc.add_paragraph()
                run = para.add_run(text)
                _set_run_style(run, 11)
                para.paragraph_format.space_after = Twips(120)
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            # Add placeholder guidance
            para = doc.add_paragraph()
            run = para.add_run(f"[This section requires development. Based on {client_type} requirements, "
                             f"this section should address the following:]")
            _set_run_style(run, 11, italic=True, color=GREY)

            # Add section-specific guidance
            guidance = _get_section_guidance(section_name, client_type)
            for bullet in guidance:
                para = doc.add_paragraph(style="List Bullet")
                run = para.add_run(bullet)
                _set_run_style(run, 11, color=GREY)

        doc.add_paragraph()  # spacing

    # ── Compliance Checklist ──
    doc.add_page_break()
    _add_heading_styled(doc, "Compliance Checklist", level=1)

    checklist = _generate_compliance_checklist(client_type, source_content)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for idx, header in enumerate(["#", "Required Section", "Status"]):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            _set_run_style(run, 10, bold=True, color=WHITE)
        from docx.oxml.ns import qn
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): "009959",
            qn("w:val"): "clear",
        })
        shading.append(shading_elm)

    for idx, item in enumerate(checklist, 1):
        row = table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = item["section"]
        row.cells[2].text = item["status"]

    # Save
    doc.save(output_path)

    # ── Summary ──
    covered = sum(1 for c in checklist if c["found"])
    total = len(checklist)

    summary = (
        f"PROPOSAL DEVELOPMENT SUMMARY\n"
        f"{'=' * 40}\n"
        f"Client Type:           {client_type}\n"
        f"Assigned Agents:       {assigned_agents}\n"
        f"Project Title:         {project_title or 'Untitled'}\n"
        f"Country/Region:        {country}\n"
        f"Source Word Count:     {source_word_count:,}\n"
        f"Required Sections:     {total}\n"
        f"Sections Populated:    {sections_populated} from source\n"
        f"Compliance Score:      {covered}/{total} ({100*covered//total if total else 0}%)\n"
        f"Sections Needing Dev:  {total - covered}\n"
        f"{'=' * 40}\n"
    )

    return summary


def _get_section_guidance(section_name, client_type):
    """Return bullet-point guidance for a specific section."""
    guidance_map = {
        "Executive Summary": [
            "Provide a concise overview of the proposed approach (max 2 pages)",
            "Highlight key differentiators and value proposition",
            "Summarize the team's relevant experience",
        ],
        "Understanding of the Assignment": [
            "Demonstrate understanding of the project context and objectives",
            "Reference the Terms of Reference (TOR) requirements",
            "Identify key challenges and constraints",
        ],
        "Technical Approach and Methodology": [
            "Describe the step-by-step technical approach",
            "Justify the chosen methodology",
            "Include data collection and analysis methods",
        ],
        "Technical Approach": [
            "Describe the step-by-step technical approach",
            "Justify the chosen methodology",
            "Include data collection and analysis methods",
        ],
        "Methodology": [
            "Detail the research/implementation methodology",
            "Include tools and frameworks to be used",
            "Describe quality assurance measures",
        ],
        "Work Plan and Schedule": [
            "Provide a detailed schedule with milestones",
            "Include Gantt chart or timeline table",
            "Identify critical path activities",
        ],
        "Work Plan": [
            "Provide a detailed schedule with milestones",
            "Include Gantt chart or timeline table",
            "Identify critical path activities",
        ],
        "Team Composition and Task Assignment": [
            "List all proposed key experts with roles",
            "Define person-month allocations per expert",
            "Include organization chart",
        ],
        "Team Composition": [
            "List key personnel with qualifications",
            "Define roles and responsibilities",
            "Include person-month allocations",
        ],
    }

    return guidance_map.get(section_name, [
        f"Develop content per {client_type} requirements",
        "Ensure compliance with submission guidelines",
        "Include supporting evidence and references",
    ])
